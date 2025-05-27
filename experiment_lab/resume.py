from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
def set_font(run, font_name="Arial", size=11):
    run.font.name = font_name
    run.font.size = Pt(size)

# Add heading (Name and Contact Info)
name = doc.add_heading("Your Name", level=1)
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in name.runs:
    set_font(run, size=14)

contact = doc.add_paragraph("Email: your.email@example.com | Phone: (123) 456-7890 | LinkedIn: linkedin.com/in/yourprofile | City, State")
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in contact.runs:
    set_font(run)

# Add horizontal line
doc.add_paragraph("_" * 80)

# Objective Section
doc.add_heading("Objective", level=2)
objective = doc.add_paragraph(
    "Motivated and detail-oriented recent graduate seeking an entry-level position in [Your Field] to apply strong technical skills and contribute to team success."
)
for run in objective.runs:
    set_font(run)

# Skills Section
doc.add_heading("Skills", level=2)
skills = [
    "Programming: Python, Java, C++",
    "Web Development: HTML, CSS, JavaScript",
    "Tools: Git, Docker, VS Code",
    "Soft Skills: Team Collaboration, Problem-Solving, Communication"
]
for skill in skills:
    p = doc.add_paragraph(skill, style="List Bullet")
    for run in p.runs:
        set_font(run)

# Projects Section
doc.add_heading("Projects", level=2)
project1 = doc.add_paragraph("Personal Portfolio Website | HTML, CSS, JavaScript | Jan 2025 - Mar 2025")
for run in project1.runs:
    set_font(run, size=11, bold=True)
p1_details = doc.add_paragraph(
    "- Developed a responsive portfolio website to showcase projects and skills.\n"
    "- Implemented features like contact forms and dynamic content using JavaScript.\n"
    "- Deployed on GitHub Pages, ensuring cross-browser compatibility."
)
p1_details.style = "List Bullet"
for run in p1_details.runs:
    set_font(run)

project2 = doc.add_paragraph("Task Management App | Python, Flask | Sep 2024 - Dec 2024")
for run in project2.runs:
    set_font(run, size=11, bold=True)
p2_details = doc.add_paragraph(
    "- Built a web application to manage tasks with user authentication.\n"
    "- Integrated SQLite database for task storage and retrieval.\n"
    "- Utilized Flask framework for backend development."
)
p2_details.style = "List Bullet"
for run in p2_details.runs:
    set_font(run)

# Education Section
doc.add_heading("Education", level=2)
edu = doc.add_paragraph("Bachelor of Science in [Your Major], [Your University], City, State | May 2025")
for run in edu.runs:
    set_font(run, size=11, bold=True)
edu_details = doc.add_paragraph("- GPA: 3.8/4.0\n- Relevant Coursework: Data Structures, Algorithms, Web Development")
edu_details.style = "List Bullet"
for run in edu_details.runs:
    set_font(run)

# Work Experience Section
doc.add_heading("Work Experience", level=2)
work = doc.add_paragraph("Intern, [Company Name], City, State | Jun 2024 - Aug 2024")
for run in work.runs:
    set_font(run, size=11, bold=True)
work_details = doc.add_paragraph(
    "- Assisted in developing internal tools using Python and SQL.\n"
    "- Collaborated with a team to streamline data processing workflows.\n"
    "- Conducted testing and debugging to ensure software reliability."
)
work_details.style = "List Bullet"
for run in work_details.runs:
    set_font(run)

# Save the document
doc.save("ATS_Friendly_Resume.docx")